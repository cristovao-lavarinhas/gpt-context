"""File text extraction for user uploads (PDF, DOCX, images via OCR)."""

from __future__ import annotations

import base64
import io
from pathlib import Path

import requests

from rag import OLLAMA_BASE_URL, OLLAMA_MODEL

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PILImage = None
    PIL_AVAILABLE = False

try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None
    TESSERACT_AVAILABLE = False


def extract_text_from_image(image_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower().lstrip(".")
    mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"

    if TESSERACT_AVAILABLE and PIL_AVAILABLE and PILImage is not None:
        try:
            img = PILImage.open(io.BytesIO(image_bytes))
            max_size = 2000
            if img.width > max_size or img.height > max_size:
                ratio = min(max_size / img.width, max_size / img.height)
                new_size = (int(img.width * ratio), int(img.height * ratio))
                img = img.resize(new_size, PILImage.Resampling.LANCZOS)
            text = pytesseract.image_to_string(img, lang="por+eng")
            if text.strip():
                return f"[OCR via Tesseract — {filename}]\n{text.strip()}"
        except Exception as e:
            print(f"Tesseract OCR failed: {e}")

    if OLLAMA_MODEL:
        try:
            img_b64 = base64.b64encode(image_bytes).decode("utf-8")
            resp = requests.post(
                f"{OLLAMA_BASE_URL}/v1/chat/completions",
                json={
                    "model": OLLAMA_MODEL,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": (
                                "Transcribe EVERY visible text in this image exactly as written. "
                                "Maintain formatting, line breaks, and structure. "
                                "Output ONLY the extracted text, nothing else."
                            )},
                            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_b64}"}},
                        ],
                    }],
                    "temperature": 0,
                    "max_tokens": 2048,
                },
                timeout=60,
            )
            if resp.status_code == 200:
                content = resp.json()["choices"][0]["message"]["content"].strip()
                if content and len(content) > 5:
                    return f"[OCR via Vision Model (Ollama) — {filename}]\n{content}"
        except Exception as e:
            print(f"Ollama vision OCR failed: {e}")

    return (
        f"⚠️ Não foi possível extrair texto da imagem '{filename}'.\n\n"
        "**Soluções:**\n"
        "1. Instala o Tesseract OCR e `pip install pytesseract`\n"
        "2. Usa um modelo Ollama com visão (ex: `ollama pull llava`) e define OLLAMA_MODEL=llava"
    )


def extract_text_from_upload(filename: str, content_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    file_size_mb = len(content_bytes) / (1024 * 1024)

    if file_size_mb > 5:
        return f"❌ Ficheiro muito grande: {file_size_mb:.1f}MB. Máximo permitido: 5MB."

    try:
        if ext == ".pdf":
            if pypdf is None:
                return "❌ Erro: pypdf não instalado. Não posso ler PDFs. Execute: `pip install pypdf`"
            pdf_reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            if not pdf_reader.pages:
                return "⚠️ PDF vazio ou corrompido."
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"[Página {i + 1}]\n{page_text}\n"
            return text if text.strip() else "⚠️ Nenhum texto extraído do PDF."

        if ext == ".docx":
            if docx is None:
                return "❌ Erro: python-docx não instalado. Não posso ler ficheiros Word. Execute: `pip install python-docx`"
            doc = docx.Document(io.BytesIO(content_bytes))
            text_parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    text_parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(text_parts) or "⚠️ Documento Word vazio."

        if ext in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".tif", ".webp"}:
            return extract_text_from_image(content_bytes, filename)

        # Plain text-like formats
        return content_bytes.decode("utf-8", errors="ignore")

    except Exception as e:
        return f"❌ Erro ao processar ficheiro '{filename}': {e}"
